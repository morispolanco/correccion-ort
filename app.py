import streamlit as st
from docx import Document
import requests # Para llamadas HTTP
import re
import io
import uuid
import time
import json # Para formatear el payload

# Configuración de la página de Streamlit
st.set_page_config(page_title="Corrector DOCX (OpenRouter)", layout="wide")

MAX_CHARS = 300000
OPENROUTER_API_URL = "https://openrouter.ai/api/v1/chat/completions"
# Puedes cambiar el modelo si lo deseas (algunos pueden tener costo en OpenRouter)
# El modelo gratuito es bueno para pruebas, pero puede tener límites más estrictos.
OPENROUTER_MODEL_NAME = "meta-llama/llama-3.3-8b-instruct:free"
# OPENROUTER_MODEL_NAME = "openai/gpt-3.5-turbo" # Ejemplo de otro modelo (puede tener costo)

# Headers recomendados por OpenRouter para identificar tu app
# Para desarrollo local, el referer puede ser un placeholder.
# Para producción, usa el dominio de tu app.
APP_SITE_URL = "http://localhost:8501" # Cambiar si se despliega
APP_TITLE = "Corrector DOCX Streamlit (Llama3)"

# --- Funciones Auxiliares (Mismas de antes para citas y conteo) ---
def extract_text_and_citations(paragraph_text):
    citations = {}
    citation_pattern = re.compile(
        r'(".*?"|\'.*?\'|\([^)]*?\d{4}[^)]*?\)|\[[^\]]*?\d{4}[^\]]*?\])'
    )
    def replace_citation(match):
        citation_text = match.group(0)
        placeholder = f"__CITATION_{uuid.uuid4().hex}__"
        citations[placeholder] = citation_text
        return placeholder
    text_with_placeholders = citation_pattern.sub(replace_citation, paragraph_text)
    return text_with_placeholders, citations

def insert_citations_back(text_with_placeholders, citations):
    final_text = text_with_placeholders
    for placeholder, original_citation in citations.items():
        final_text = final_text.replace(placeholder, original_citation, 1)
    return final_text

def get_total_characters(doc):
    count = 0
    for para in doc.paragraphs:
        count += len(para.text)
    return count

# --- Funciones para OpenRouter ---

@st.cache_data(ttl=3600) # Cache para evitar llamadas repetidas
def correct_text_with_openrouter(text_to_correct, api_key_to_use, retries=3, delay=5):
    if not text_to_correct.strip():
        return ""
    if not api_key_to_use:
        st.error("API Key de OpenRouter no disponible para la corrección.")
        return text_to_correct

    headers = {
        "Authorization": f"Bearer {api_key_to_use}",
        "Content-Type": "application/json",
        "HTTP-Referer": APP_SITE_URL, # Opcional, pero recomendado por OpenRouter
        "X-Title": APP_TITLE        # Opcional, pero recomendado
    }

    # Prompt para el modelo de chat
    system_prompt = """Eres un asistente experto en gramática y ortografía del idioma español.
Tu tarea es corregir el texto que te proporcionará el usuario.
IMPORTANTE:
1. Corrige únicamente errores gramaticales y ortográficos.
2. NO cambies el significado original del texto.
3. NO alteres ni modifiques las citas textuales que están marcadas con placeholders como __CITATION_HEXADECIMAL__. Debes dejarlas exactamente como están. Por ejemplo, si ves "__CITATION_a1b2c3d4__", esa cadena debe permanecer idéntica en tu respuesta.
4. Devuelve SOLAMENTE el texto corregido, sin ninguna introducción, explicación, saludo, despedida o comentario adicional. No escribas "Texto corregido:" ni nada similar. Solo el texto.
"""
    payload = {
        "model": OPENROUTER_MODEL_NAME,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text_to_correct}
        ],
        "temperature": 0.2, # Más bajo para ser más determinista
        # "max_tokens": 1024 # Puedes ajustar esto si es necesario
    }

    for attempt in range(retries):
        try:
            response = requests.post(OPENROUTER_API_URL, headers=headers, data=json.dumps(payload), timeout=90) # Timeout más largo
            response.raise_for_status() # Lanza una excepción para códigos de error HTTP (4xx o 5xx)
            
            response_json = response.json()
            
            if response_json.get("choices") and len(response_json["choices"]) > 0:
                corrected_text = response_json["choices"][0].get("message", {}).get("content")
                if corrected_text:
                    # Limpieza adicional de placeholders
                    for placeholder in re.findall(r"__CITATION_[a-f0-9]{32}__", corrected_text):
                        corrected_text = corrected_text.replace(f" {placeholder} ", placeholder)
                        corrected_text = corrected_text.replace(f" {placeholder}", placeholder)
                        corrected_text = corrected_text.replace(f"{placeholder} ", placeholder)
                    return corrected_text.strip()
            
            # Si no hay 'choices' o 'content', o la respuesta es inesperada
            st.warning(f"Intento {attempt + 1} de {retries}: OpenRouter devolvió una respuesta inesperada o vacía. Detalle: {response_json.get('error', 'Sin detalle de error')}")
            time.sleep(delay * (attempt + 1)) # Backoff exponencial

        except requests.exceptions.HTTPError as e:
            error_content = e.response.json() if e.response else {}
            error_message = error_content.get('error', {}).get('message', str(e))
            st.warning(f"Error HTTP de OpenRouter (intento {attempt + 1}/{retries}): {e.response.status_code} - {error_message}")
            if e.response.status_code == 401: # Unauthorized
                st.error("Error: La API Key de OpenRouter proporcionada no es válida o no tiene permisos. Por favor, verifica.")
                return None # Error fatal, no reintentar
            if e.response.status_code == 429: # Rate limit
                 st.warning("Límite de tasa alcanzado. Esperando más tiempo...")
                 time.sleep(delay * (attempt + 1) * 2) # Esperar más en caso de rate limit
            elif attempt == retries - 1:
                st.error(f"No se pudo corregir el fragmento después de {retries} intentos: {text_to_correct[:100]}...")
                return text_to_correct
            else:
                time.sleep(delay * (attempt + 1))
        except requests.exceptions.RequestException as e:
            st.warning(f"Error de red o conexión con OpenRouter (intento {attempt + 1}/{retries}): {e}")
            if attempt == retries - 1:
                st.error(f"No se pudo conectar con OpenRouter para el fragmento: {text_to_correct[:100]}...")
                return text_to_correct
            time.sleep(delay * (attempt + 1))
        except Exception as e: # Captura general para otros errores inesperados
            st.error(f"Error inesperado al procesar con OpenRouter (intento {attempt + 1}/{retries}): {e}")
            if attempt == retries - 1:
                return text_to_correct # Devolver original si todos los reintentos fallan
            time.sleep(delay * (attempt + 1))

    st.warning(f"Fragmento no procesado por OpenRouter después de {retries} intentos, se mantendrá original: {text_to_correct[:100]}...")
    return text_to_correct

def process_document_openrouter(doc_bytes, api_key_to_use):
    if not api_key_to_use:
        st.error("Por favor, configura tu API Key de OpenRouter para continuar.")
        return None

    doc = Document(io.BytesIO(doc_bytes))
    total_chars = get_total_characters(doc)
    if total_chars > MAX_CHARS:
        st.error(f"El documento excede el límite de {MAX_CHARS} caracteres (tiene {total_chars}).")
        return None

    new_doc = Document()
    for style in doc.styles:
        try: new_doc.styles.add_style(style.name, style.type)
        except ValueError: pass

    progress_bar = st.progress(0)
    total_paragraphs = len(doc.paragraphs)
    processed_paragraphs = 0
    # Estimar el tiempo total puede ser difícil, pero podemos mostrar el progreso por párrafos
    st.info(f"Procesando aproximadamente {total_paragraphs} párrafos. El modelo '{OPENROUTER_MODEL_NAME}' puede ser lento.")

    for para_idx, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            new_para = new_doc.add_paragraph()
            if para.style and para.style.name in new_doc.styles:
                 new_para.style = new_doc.styles[para.style.name]
            new_para.paragraph_format.alignment = para.paragraph_format.alignment
            processed_paragraphs += 1
            progress_bar.progress(processed_paragraphs / total_paragraphs)
            continue

        original_text = para.text
        text_with_placeholders, citations = extract_text_and_citations(original_text)
        
        if text_with_placeholders.strip():
            corrected_text_with_placeholders = correct_text_with_openrouter(text_with_placeholders, api_key_to_use)
            if corrected_text_with_placeholders is None: # Error fatal con API Key
                return None 
        else:
            corrected_text_with_placeholders = text_with_placeholders

        final_corrected_text = insert_citations_back(corrected_text_with_placeholders, citations)
        new_para = new_doc.add_paragraph()

        if para.style and para.style.name in new_doc.styles:
            new_para.style = new_doc.styles[para.style.name]
        else:
            new_para.style = new_doc.styles['Normal']

        # Copiar formato de párrafo
        new_para.paragraph_format.alignment = para.paragraph_format.alignment
        new_para.paragraph_format.left_indent = para.paragraph_format.left_indent
        new_para.paragraph_format.right_indent = para.paragraph_format.right_indent
        new_para.paragraph_format.first_line_indent = para.paragraph_format.first_line_indent
        new_para.paragraph_format.space_before = para.paragraph_format.space_before
        new_para.paragraph_format.space_after = para.paragraph_format.space_after
        new_para.paragraph_format.line_spacing = para.paragraph_format.line_spacing

        if para.runs:
            first_run_style = para.runs[0]
            run = new_para.add_run(final_corrected_text)
            run.bold = first_run_style.bold
            run.italic = first_run_style.italic
            run.underline = first_run_style.underline
            if first_run_style.font.name: run.font.name = first_run_style.font.name
            if first_run_style.font.size: run.font.size = first_run_style.font.size
            if first_run_style.font.color and first_run_style.font.color.rgb:
                 run.font.color.rgb = first_run_style.font.color.rgb
        else:
            new_para.add_run(final_corrected_text)

        processed_paragraphs += 1
        progress_bar.progress(processed_paragraphs / total_paragraphs)
        # Pausa más larga si se usa un modelo gratuito con rate limits estrictos
        # O si la API es simplemente más lenta.
        time.sleep(0.2 if "free" in OPENROUTER_MODEL_NAME.lower() else 0.1)


    doc_buffer = io.BytesIO()
    new_doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

# --- Interfaz de Streamlit ---
st.title(f"✍️ Corrector DOCX (con OpenRouter: {OPENROUTER_MODEL_NAME})")
st.markdown(f"""
Sube un archivo Word (.docx) de hasta **{MAX_CHARS // 1000}k caracteres**.
La aplicación corregirá la gramática y ortografía usando la API de OpenRouter.ai
con el modelo **{OPENROUTER_MODEL_NAME}**.
""")

st.sidebar.header("Configuración de API Key (OpenRouter)")
api_key_from_secrets = ""
try:
    if "OPENROUTER_API_KEY" in st.secrets and st.secrets["OPENROUTER_API_KEY"]:
        api_key_from_secrets = st.secrets["OPENROUTER_API_KEY"]
        st.sidebar.success("API Key de OpenRouter cargada desde secretos.")
    else:
        st.sidebar.info("API Key de OpenRouter no encontrada en `secrets.toml` o está vacía.")
except (FileNotFoundError, AttributeError):
    st.sidebar.info("Archivo `secrets.toml` no encontrado. Ingresa la API Key manualmente.")

user_provided_api_key = st.sidebar.text_input(
    "Ingresa tu OpenRouter API Key (opcional si está en secrets.toml)",
    type="password",
    value=api_key_from_secrets,
    help="Obtén tu API key en OpenRouter.ai. Si está en secrets.toml, se usará esa por defecto."
)

final_api_key_to_use = user_provided_api_key
if not user_provided_api_key and api_key_from_secrets:
    final_api_key_to_use = api_key_from_secrets

uploaded_file = st.file_uploader("Carga tu archivo .docx aquí", type="docx")

if uploaded_file is not None:
    file_bytes = uploaded_file.getvalue()
    st.info(f"Archivo cargado: {uploaded_file.name} ({len(file_bytes) / 1024:.2f} KB)")

    if st.button(f"🔎 Corregir Documento con OpenRouter ({OPENROUTER_MODEL_NAME.split('/')[-1].split(':')[0]})"):
        if not final_api_key_to_use:
            st.error("Error: Falta la API Key de OpenRouter. Por favor, ingrésala o configúrala en `secrets.toml`.")
        else:
            with st.spinner(f"Procesando con OpenRouter ({OPENROUTER_MODEL_NAME})... Esto puede tardar."):
                try:
                    corrected_doc_buffer = process_document_openrouter(file_bytes, final_api_key_to_use)
                    if corrected_doc_buffer:
                        st.success("¡Documento procesado y corregido con éxito!")
                        st.download_button(
                            label="📥 Descargar Documento Corregido (.docx)",
                            data=corrected_doc_buffer,
                            file_name=f"corregido_openrouter_{uploaded_file.name}",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Ocurrió un error durante el procesamiento con OpenRouter: {e}")
                    st.exception(e)
else:
    st.info("Esperando a que se suba un archivo .docx.")

st.markdown("---")
st.markdown(f"Desarrollado con Streamlit, python-docx y OpenRouter.ai ({OPENROUTER_MODEL_NAME}).")
